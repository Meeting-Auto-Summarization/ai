from flask import Flask, request, jsonify, Response
from flask_cors import CORS
import torch
from transformers import PreTrainedTokenizerFast
from transformers import BartForConditionalGeneration
import json
from io import BytesIO
from docx import Document
import math

app = Flask(__name__)

CORS(app, resources={r'/*': {'origins': '*'}})

tokenizer = PreTrainedTokenizerFast.from_pretrained('gogamza/kobart-base-v1')
model = BartForConditionalGeneration.from_pretrained('kijun/mas-kobart-v1')

@app.route('/summarize', methods=['POST', 'GET'])
def hello_world():
    if request.method == 'POST':
        text = json.loads(request.get_data())['contents']
        return jsonify(predict(text))
    else:
        return jsonify('Invalid Method')

@app.route('/script-docx', methods=['POST', 'GET'])
def get_script_docx():
    if request.method == 'POST':
        data = json.loads(request.get_data())
        meeting = data['meeting']
        script = data['script']

        output_stream = BytesIO()

        document = Document()

        document.add_heading(meeting['title'], 0)

        table = document.add_table(rows=len(script), cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Time'
        hdr_cells[2].text = 'Content'
        for i in range(len(script)):
            row_cells = table.add_row().cells
            row_cells[0].text = script[i]['nick']

            stime = int(script[i]['time'])

            hour = int(stime / 3600)
            min = int((stime / 60) % 60)
            sec = int(stime % 60)

            if hour == 0:
                stime = f'{min}:{sec}'
            else:
                stime = f'{hour}:{min}:{sec}'
            
            row_cells[1].text = stime
            row_cells[2].text = script[i]['content']

        document.save(output_stream)

        title = meeting['title']
        filename = f'{title}_script.docx'.encode('utf-8')

        response = Response(
            output_stream.getvalue(), 
            mimetype='application/docx',
            content_type='application/octet-stream'
        )
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        output_stream.close()

        return response
    else:
        return jsonify('Invalid Method')

@app.route('/report-docx', methods=['POST', 'GET'])
def get_report_docx():
    if request.method == 'POST':
        data = json.loads(request.get_data())
        meeting = data['meeting']
        report = data['report']

        output_stream = BytesIO()

        document = Document()

        document.add_heading(meeting['title'], 0)

        for i in range(len(report)):
            for j in range(len(report[i])):
                if j == 0:
                    document.add_heading(report[i][j]['title'], level=1)
                else:
                    heading = document.add_heading(report[i][j]['title'], level=2)
                    p = document.add_paragraph(report[i][j]['summary'], style='List Bullet')

        document.save(output_stream)

        title = meeting['title']
        filename = f'{title}_report.docx'.encode('utf-8')

        response = Response(
            output_stream.getvalue(), 
            mimetype='application/docx',
            content_type='application/octet-stream',
        )
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        output_stream.close()

        return response
    else:
        return jsonify('Invalid Method')

@app.route('/script-txt', methods=['POST', 'GET'])
def get_script_txt():
    if request.method == 'POST':
        data = json.loads(request.get_data())
        meeting = data['meeting']
        script = data['script']

        title = meeting['title']
        date = meeting['date']
        members = ', '.join(meeting['members'])

        output_stream = BytesIO()

        intro = f'회의 제목: {title}\n회의 일시: {date}\n참여 인원: {members}\n\nname\ttime\tcontent\n'.encode('utf-8')
        output_stream.write(intro)

        for line in script:
            time = int(line['time'])
            nick = line['nick']
            content = line['content']

            seconds = int(time % 60)
            minutes = int((time / 60) % 60)
            hours = int(time / 3600)

            timeStr = f'{hours}:{str(0) + str(minutes) if minutes < 10 else minutes}:{str(0) + str(seconds) if seconds < 10 else seconds}'

            text = f'{nick}\t{timeStr}\t{content}\n'.encode('utf-8')
            output_stream.write(text)

        filename = f'{title}_script.txt'.encode('utf-8')

        response = Response(
            output_stream.getvalue(), 
            mimetype='text/txt',
            content_type='application/octet-stream',
        )
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        output_stream.close()

        return response
    else:
        return jsonify('Invalid Method')

@app.route('/report-txt', methods=['POST', 'GET'])
def get_report_txt():
    if request.method == 'POST':
        data = json.loads(request.get_data())
        meeting = data['meeting']
        report = data['report']

        title = meeting['title']
        date = meeting['date']
        members = ', '.join(meeting['members'])

        output_stream = BytesIO()

        intro = f'회의 제목: {title}\n회의 일시: {date}\n참여 인원: {members}\n\n'.encode('utf-8')
        output_stream.write(intro)

        for index, onedim in enumerate(report):
            for subIndex, reportItem in enumerate(onedim):
                report_title = reportItem['title']
                summary = reportItem['summary']

                if subIndex == 0:
                    output_stream.write(f'{index + 1}. {report_title}\n'.encode('utf-8'))
                    if len(onedim) == 1:
                        output_stream.write(f'\t{summary}\n'.encode('utf-8'))
                else:
                    output_stream.write(f'\t{chr(subIndex + 96)}. {report_title}\n\t\t{summary}\n'.encode('utf-8'))

        filename = f'{title}_report.txt'.encode('utf-8')

        response = Response(
            output_stream.getvalue(), 
            mimetype='text/txt',
            content_type='application/octet-stream',
        )
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        output_stream.close()

        return response
    else:
        return jsonify('Invalid Method')

def predict(text):
    summaryList = []

    for i in range(len(text)):
        summaryList.append([])
        for j in range(len(text[i])):
            if text[i][j] != '':
                input_ids = tokenizer.encode(text[i][j])
                input_ids = [tokenizer.bos_token_id] + input_ids + [tokenizer.eos_token_id]
                input_ids = torch.tensor(input_ids)
                input_ids = input_ids.unsqueeze(0)
                size = input_ids.size(1)
                print(size)

                if size > 150:
                    tensor_size = math.ceil(size / 150)
                    sub_input_ids = torch.chunk(input_ids, tensor_size, 1)
                    label = ''
                    for k in range(tensor_size):
                        output = model.generate(sub_input_ids[k], eos_token_id=1, max_length=512, num_beams=5)
                        output = tokenizer.decode(output[0], skip_special_tokens=True)
                        label += f'{output}\n'
                    summaryList[i].append(label)
                else:
                    output = model.generate(input_ids, eos_token_id=1, max_length=512, num_beams=5)
                    output = tokenizer.decode(output[0], skip_special_tokens=True)
                    print(output)
                    summaryList[i].append(output)
            else:
                summaryList[i].append('')

    return summaryList


if __name__ == '__main__':
    app.run(debug=True)
