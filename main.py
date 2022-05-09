from fastapi import FastAPI, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import numpy as np
from transformers import PreTrainedTokenizerFast, BartForConditionalGeneration, pipeline
from sentence_transformers import SentenceTransformer, util
from io import BytesIO
from docx import Document
from kiwipiepy import Kiwi
import ray, os

os.environ["TOKENIZERS_PARALLELISM"] = "true"

class Item(BaseModel):
    contents: list

class ReportItem(BaseModel):
    meeting: object
    report: list

class ScriptItem(BaseModel):
    meeting: object
    script: list

tokenizer = PreTrainedTokenizerFast.from_pretrained('gogamza/kobart-base-v1')
model = BartForConditionalGeneration.from_pretrained('kijun/mas-kobart-v1')
embedder = SentenceTransformer('jhgan/ko-sbert-sts')

summarizer = pipeline("summarization", model=model, tokenizer=tokenizer, framework="pt")

@ray.remote
def predict(pipeline, text_data):
    return pipeline(text_data, max_length=512)[0]['summary_text']

app = FastAPI()
origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"Hello": "World"}

@app.post("/summarize")
def summarize(body: Item):
    text = body.contents

    summaryList = []

    for i in range(len(text)):
        summaryList.append([])
        for j in range(len(text[i])):
            if text[i][j] != '':
                segment_text = segmentation(text[i][j])

                # ray.init(num_cpus=4, ignore_reinit_error=True)
                # pipe_id = ray.put(summarizer)
                # predictions = ray.get([predict.remote(pipe_id, corpus) for corpus in segment_text])
                # ray.shutdown()
                predictions = summarizer(segment_text, max_length=512)
                label = '\n'.join([text['summary_text'] for text in predictions])

                # label = '\n'.join(predictions)
                summaryList[i].append(label)
            else:
                summaryList[i].append('')

    return summaryList

def segmentation(corpus):
    kiwi = Kiwi()
    split_sentences = kiwi.split_into_sents(corpus)
    # corpus = corpus.replace("\\", "").split(".")
    # corpus = [v.strip() for v in corpus if v]

    split_corpus = []

    flag = False
    for i in range(0, len(split_sentences) - 1):
        if flag:
            flag = False
        else:
            if split_sentences[i].text[-2].isdigit() and split_sentences[i + 1].text[0].isdigit():
                split_corpus.append(f"{split_sentences[i].text}{split_sentences[i + 1].text}")
                flag = True
            else:
                split_corpus.append(split_sentences[i].text)
    
    if len(split_corpus) < 6:
        return corpus

    corpus_embeddings = embedder.encode(split_corpus, convert_to_tensor=True)
    corpus_embeddings.size()

    similarity_timeseries = []

    for index in range(corpus_embeddings.size()[0] - 1):
        similarity = float(util.pytorch_cos_sim(corpus_embeddings[index], corpus_embeddings[index + 1]))
        similarity_timeseries.append(similarity)

    similarity_timeseries = np.array(similarity_timeseries)
    threshold = similarity_timeseries.mean() - similarity_timeseries.var() + 0.03

    segment_index = np.where(similarity_timeseries < threshold)[0] + 1

    segment_corpus = []

    prev_index = 0

    for index in segment_index:
        corpus_list = ' '.join(split_corpus[prev_index:index])
        prev_index = index
        segment_corpus.append(corpus_list)

    if prev_index != len(split_corpus):
        corpus_list = ' '.join(split_corpus[prev_index:len(split_corpus)])
        segment_corpus.append(corpus_list)

    print("segment list start")
    for index, segment in enumerate(segment_corpus):
        print(f"{index}. {segment}")
    print("\n")

    return segment_corpus

@app.post("/script-docx")
def get_script_docx(body: ScriptItem):
    meeting = body.meeting
    script = body.script

    output_stream = BytesIO()

    document = Document()

    document.add_heading(meeting['title'], 0)

    table = document.add_table(rows=len(script), cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Time'
    hdr_cells[2].text = 'Content'
    for i in range(len(script)):
        row_cells = table.add_row().cells
        row_cells[0].text = script[i]['nick']

        stime = int(script[i]['time'])

        hour = int(stime / 3600)
        min = int((stime / 60) % 60)
        sec = int(stime % 60)

        if hour == 0:
            stime = f'{min}:{sec}'
        else:
            stime = f'{hour}:{min}:{sec}'
        
        row_cells[1].text = stime
        row_cells[2].text = script[i]['content']

    document.save(output_stream)

    title = meeting['title']
    filename = f'{title}_script.docx'.encode('utf-8')

    headers = {
        "mimetype": "application/docx",
        "content-type": "application/octet-stream",
        "content-disposition": f"attachment; filename={filename}"
    }

    content = output_stream.getvalue()
    output_stream.close()

    return Response(content=content, headers=headers)

@app.post("/report-docx")
def get_script_docx(body: ReportItem):
    meeting = body.meeting
    report = body.report
    
    output_stream = BytesIO()

    document = Document()

    document.add_heading(meeting['title'], 0)

    for i in range(len(report)):
        for j in range(len(report[i])):
            if j == 0:
                document.add_heading(report[i][j]['title'], level=1)
            else:
                heading = document.add_heading(report[i][j]['title'], level=2)
                p = document.add_paragraph(report[i][j]['summary'], style='List Bullet')

    document.save(output_stream)

    title = meeting['title']
    filename = f'{title}_report.docx'.encode('utf-8')

    headers = {
        "mimetype": "application/docx",
        "content-type": "application/octet-stream",
        "content-disposition": f"attachment; filename={filename}"
    }

    content = output_stream.getvalue()
    output_stream.close()

    return Response(content=content, headers=headers)

@app.post("/script-txt")
def get_script_txt(body: ScriptItem):
    meeting = body.meeting
    script = body.script

    title = meeting['title']
    date = meeting['date']
    members = ', '.join(meeting['members'])

    output_stream = BytesIO()

    intro = f'회의 제목: {title}\n회의 일시: {date}\n참여 인원: {members}\n\nname\ttime\tcontent\n'.encode('utf-8')
    output_stream.write(intro)

    for line in script:
        time = int(line['time'])
        nick = line['nick']
        content = line['content']

        seconds = int(time % 60)
        minutes = int((time / 60) % 60)
        hours = int(time / 3600)

        timeStr = f'{hours}:{str(0) + str(minutes) if minutes < 10 else minutes}:{str(0) + str(seconds) if seconds < 10 else seconds}'

        text = f'{nick}\t{timeStr}\t{content}\n'.encode('utf-8')
        output_stream.write(text)

    filename = f'{title}_script.txt'.encode('utf-8')

    headers = {
        "mimetype": "application/docx",
        "content-type": "application/octet-stream",
        "content-disposition": f"attachment; filename={filename}"
    }

    content = output_stream.getvalue()
    output_stream.close()

    return Response(content=content, headers=headers)


@app.post("/report-txt")
def get_report_txt(body: ReportItem):
    meeting = body.meeting
    report = body.report

    title = meeting['title']
    date = meeting['date']
    members = ', '.join(meeting['members'])

    output_stream = BytesIO()

    intro = f'회의 제목: {title}\n회의 일시: {date}\n참여 인원: {members}\n\n'.encode('utf-8')
    output_stream.write(intro)

    for index, onedim in enumerate(report):
        for subIndex, reportItem in enumerate(onedim):
            report_title = reportItem['title']
            summary = reportItem['summary']

            if subIndex == 0:
                output_stream.write(f'{index + 1}. {report_title}\n'.encode('utf-8'))
                if len(onedim) == 1:
                    output_stream.write(f'\t{summary}\n'.encode('utf-8'))
            else:
                output_stream.write(f'\t{chr(subIndex + 96)}. {report_title}\n\t\t{summary}\n'.encode('utf-8'))

    filename = f'{title}_report.txt'.encode('utf-8')

    headers = {
        "mimetype": "application/docx",
        "content-type": "application/octet-stream",
        "content-disposition": f"attachment; filename={filename}"
    }

    content = output_stream.getvalue()
    output_stream.close()

    return Response(content=content, headers=headers)

